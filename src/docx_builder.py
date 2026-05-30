"""Low-level helpers for building the Word report from the BIHSO template.

The strategy is to reuse ``BIHSO Projekt.docx`` as a template so the new
document inherits an identical title page and styling (Aptos fonts, heading
styles, A4 margins, footer page numbers). We:

1. open the template,
2. replace the work title on the title page,
3. strip the old body content (everything from the "Streszczenie" heading down
   to, but excluding, the final section properties),
4. append fresh content using the template's own paragraph/table styles.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

# Style names exactly as exposed by python-docx for this template.
STYLE_H2 = "Heading 2"
STYLE_H3 = "Heading 3"
STYLE_NORMAL = "Normal"
STYLE_CAPTION = "Caption"
STYLE_LIST = "List Paragraph"
STYLE_TABLE = "Table Grid"

# Body child index of the title paragraph on the title page (verified against
# the template structure).
TITLE_PARAGRAPH_INDEX = 12
# Heading text that marks the start of the removable body content.
CONTENT_START_HEADING = "Streszczenie"


class ReportBuilder:
    """Builds the report on top of the BIHSO template."""

    def __init__(self, template_path: Path) -> None:
        self.doc = Document(str(template_path))
        self.body = self.doc.element.body
        self._sectPr = None

    # --- template preparation ---------------------------------------------
    def set_work_title(self, title: str) -> None:
        """Replace the title-page work title, keeping its run formatting."""
        para = Paragraph(self.body[TITLE_PARAGRAPH_INDEX], self.doc)
        runs = para.runs
        if not runs:
            para.add_run(title)
            return
        # Put the whole title in the first run, clear the remaining ones so the
        # original bold/size formatting (14 pt bold, centered) is preserved.
        runs[0].text = title
        for r in runs[1:]:
            r.text = ""

    def strip_old_body(self) -> None:
        """Remove every body element after the title page, keeping sectPr."""
        children = list(self.body)
        # Locate the final section properties to preserve page setup + footers.
        self._sectPr = None
        if children and children[-1].tag == qn("w:sectPr"):
            self._sectPr = children[-1]

        start = None
        for i, ch in enumerate(children):
            if ch.tag == qn("w:p"):
                p = Paragraph(ch, self.doc)
                if (p.style is not None
                        and p.style.name == STYLE_H2
                        and p.text.strip() == CONTENT_START_HEADING):
                    start = i
                    break
        if start is None:
            raise RuntimeError("Could not locate the content start heading.")

        for ch in children[start:]:
            if ch is self._sectPr:
                continue
            self.body.remove(ch)
        # Note: sectPr is intentionally kept in the body so that python-docx can
        # resolve the current section (needed by add_table). It is moved to the
        # very end in finalize(), after the new content has been appended.

    def finalize(self, out_path: Path) -> None:
        """Move section properties to the very end of the body and save."""
        if self._sectPr is not None:
            self.body.remove(self._sectPr)
            self.body.append(self._sectPr)
        self.doc.save(str(out_path))

    # --- content helpers ---------------------------------------------------
    def h2(self, text: str) -> Paragraph:
        return self.doc.add_paragraph(text, style=STYLE_H2)

    def h3(self, text: str) -> Paragraph:
        return self.doc.add_paragraph(text, style=STYLE_H3)

    def para(self, text: str, justify: bool = True) -> Paragraph:
        p = self.doc.add_paragraph(text, style=STYLE_NORMAL)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    def bullets(self, items: list[str]) -> None:
        for it in items:
            p = self.doc.add_paragraph(style=STYLE_LIST)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(0.8)
            self._add_bullet_run(p, it)

    @staticmethod
    def _add_bullet_run(p: Paragraph, text: str) -> None:
        # Render a simple manual bullet so we do not depend on list numbering.
        run = p.add_run("\u2022  " + text)

    def caption(self, text: str) -> Paragraph:
        p = self.doc.add_paragraph(text, style=STYLE_CAPTION)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def figure(self, image_path: Path, caption: str, width_cm: float = 15.0) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(width_cm))
        self.caption(caption)

    def spacer(self) -> None:
        self.doc.add_paragraph("", style=STYLE_NORMAL)

    def code_block(self, code: str, caption: str | None = None) -> None:
        """Render a monospaced, lightly shaded code listing."""
        p = self.doc.add_paragraph(style=STYLE_NORMAL)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.4)
        pf.space_after = Pt(2)
        pf.space_before = Pt(2)
        run = p.add_run(code.strip("\n"))
        run.font.name = "Consolas"
        # Ensure the east-asian/complex-script font also maps to a monospace.
        run._element.rPr.rFonts.set(qn("w:cs"), "Consolas")
        run.font.size = Pt(9)
        self._shade_paragraph(p, "F2F2F2")
        if caption:
            self.caption(caption)

    @staticmethod
    def _shade_paragraph(p: Paragraph, hex_fill: str) -> None:
        """Apply a background shading to a paragraph (light grey by default)."""
        pPr = p._p.get_or_add_pPr()
        shd = pPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): hex_fill,
        })
        pPr.append(shd)

    def table(self, headers: list[str], rows: list[list[str]],
              caption: str | None = None, font_pt: float = 9.5,
              col_widths_cm: list[float] | None = None) -> None:
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.style = STYLE_TABLE
        tbl.alignment = 1  # center

        hdr = tbl.rows[0].cells
        for i, h in enumerate(headers):
            self._set_cell(hdr[i], h, bold=True, font_pt=font_pt)
        for row in rows:
            cells = tbl.add_row().cells
            for i, val in enumerate(row):
                self._set_cell(cells[i], val, bold=False, font_pt=font_pt)

        if col_widths_cm:
            for i, w in enumerate(col_widths_cm):
                for row in tbl.rows:
                    row.cells[i].width = Cm(w)

        if caption:
            self.caption(caption)

    @staticmethod
    def _set_cell(cell, text: str, bold: bool, font_pt: float) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_pt)
        if bold:
            # Light shading effect via dark text; keep it simple and robust.
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
